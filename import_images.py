"""
Import hình ảnh từ file Word (.docx) vào database cho sản phẩm Két Nhớt & Két Nước.

Cách dùng:
    py import_images.py --dry-run              # Chỉ preview, ko ghi DB
    py import_images.py --limit 2              # Demo: chỉ import 2 SP mỗi file
    py import_images.py --create-missing       # Tạo SP mới nếu chưa có trong DB
    py import_images.py --create-missing --dry-run  # Preview tạo mới trước
    py import_images.py                        # Import tất cả (chỉ update SP có sẵn)
"""

import sys
import os
import re
import io
import argparse
from pathlib import Path
from decimal import Decimal

# Fix encoding for Windows terminals
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# Django setup
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
env_file = Path(__file__).resolve().parent / 'backend' / '.env'
if env_file.exists():
    for line in env_file.read_text(encoding='utf-8').splitlines():
        line = line.strip()
        if line and not line.startswith('#') and '=' in line:
            key, _, val = line.partition('=')
            os.environ.setdefault(key.strip(), val.strip())

import django
django.setup()

from django.db.models import Q, Max
from products.models import Product, HangMay
from docx import Document


# ── Config ──
DOCS_DIR = Path(__file__).resolve().parent / 'docs'
MEDIA_DIR = Path(__file__).resolve().parent / 'media' / 'products'

FILE_LOAI_MAP = {
    'KET NHOT': 'ket_nhot',
    'KET NUOC': 'ket_nuoc',
}

# Mã model regex MỞ RỘNG
MODEL_PATTERNS = [
    # Komatsu
    r'PC\d+[-/]\d+', r'PC\d+[A-Z]?', r'PC\d+US[-/]\d+',
    # Hitachi EX
    r'EX\d+[-/]\d+', r'EX\d+[A-Z]?',
    # Hitachi ZX
    r'ZX\d+[-/]\d+', r'ZX\d+[A-Z]?',
    # Kobelco SK
    r'SK\d+[-/]\d+', r'SK\d+[A-Z]?', r'SK\d+N\d+',
    # Sumitomo SH
    r'SH\d+[A-Z]?\d*', r'SH\d+[-/]\d+',
    # Hyundai R
    r'R\d+[-/]\d+', r'R\d+[A-Z]?',
    # Doosan/Daewoo DH, DX
    r'DH\d+[-/]\d+', r'DH\d+[A-Z]?',
    r'DX\d+[-/]\d+', r'DX\d+[A-Z]?',
    # Volvo EC
    r'EC\d+[A-Z]?\d*', r'EC\d+[-/]\d+',
    # Caterpillar E
    r'E\d+[A-Z]\d*', r'E\d+[A-Z]',
    # Engine codes
    r'S\d+[A-Z]\d+',
    # Other
    r'D\d+[A-Z]?[-/]\d+', r'D\d+[A-Z]?',
    r'HD\d+[-/]\d+',
    r'WA\d+[-/]\d+', r'WA\d+',
    # Yanmar, IHI
    r'YANMAR\s*\d+', r'IHI\d+',
    # ZX without dash
    r'ZX\d+',
]


def extract_model_codes(name: str) -> list[str]:
    """Trích xuất các mã model từ tên sản phẩm trong docx."""
    codes = []
    for pattern in MODEL_PATTERNS:
        matches = re.findall(pattern, name, re.IGNORECASE)
        codes.extend(matches)
    # Chuẩn hóa, loại bỏ trùng
    seen = set()
    result = []
    for c in codes:
        c_clean = c.upper().replace(' ', '')
        if c_clean not in seen:
            seen.add(c_clean)
            result.append(c_clean)
    return result


def detect_brand(name: str) -> str:
    """Nhận diện hãng máy từ tên sản phẩm."""
    n = name.upper()
    if 'HYUNDAI' in n or re.search(r'\bR\d{2,3}', n):
        return 'HYUNDAI'
    if 'VOLVO' in n or re.search(r'\bEC\d{2,3}', n):
        return 'VOLVO'
    if 'DOOSAN' in n or 'DAEWOO' in n or re.search(r'\bD[HX]\d{2,3}', n):
        return 'DOOSAN/DAEWOO'
    if 'SUMITOMO' in n or re.search(r'\bSH\d{2,4}A?\d*', n):
        return 'SUMITOMO'
    if 'KOBELCO' in n or re.search(r'\bSK\d{2,4}', n):
        return 'KOBELCO'
    if 'HITACHI' in n or re.search(r'\bZX\d{2,4}', n) or re.search(r'\bEX\d{2,4}', n):
        return 'HITACHI'
    if 'KOMATSU' in n or 'PC' in n or re.search(r'\bPC\d{2,4}', n):
        return 'KOMATSU'
    if 'CATERPILLAR' in n or 'CAT' in n or re.search(r'\bE\d{3,4}[A-Z]', n):
        return 'CATERPILLAR'
    if 'IHI' in n:
        return 'IHI'
    if 'YANMAR' in n:
        return 'YANMAR'
    return 'CHƯA RÕ'


# HangMay lookup cache
_HANGMAY_MAP = {
    'KOMATSU': 'KOMATSU',
    'HITACHI': 'HITACHI',
    'KOBELCO': 'KOBELCO',
    'CATERPILLAR': 'CATERPILLAR',
    'HYUNDAI': 'HYUNDAI',
    'VOLVO': 'VOLVO',
    'DOOSAN/DAEWOO': 'DOOSAN/DAEWOO',
    'SUMITOMO': 'SUMITOMO',
    'IHI': 'IHI',
    'YANMAR': 'YANMAR',
    'CHƯA RÕ': 'CHƯA RÕ',
}


def get_hang_may(brand_name: str) -> HangMay | None:
    """Tìm hoặc tạo HangMay từ tên brand."""
    search_name = _HANGMAY_MAP.get(brand_name, brand_name)
    hm = HangMay.objects.filter(ten__iexact=search_name).first()
    if hm:
        return hm
    # Fallback: tìm gần đúng
    hm = HangMay.objects.filter(ten__icontains=search_name).first()
    if hm:
        return hm
    # Tạo mới nếu không có
    from django.utils.text import slugify
    hm = HangMay.objects.create(
        ten=search_name,
        slug=slugify(search_name),
    )
    return hm


def get_next_ma_vt() -> str:
    """Sinh mã VT mới dạng HHxxxxx."""
    last = Product.objects.filter(ma_vt__regex=r'^HH\d+$').aggregate(m=Max('ma_vt'))['m']
    if last and last[2:].isdigit():
        next_num = int(last[2:]) + 1
    else:
        next_num = 90000
    return f'HH{next_num}'


def parse_docx(filepath: Path) -> list[dict]:
    """
    Parse file Word, trả về list các dict:
    [{name: 'KÉT NHỚT PC60-5 / PC60-6', image_blobs: [bytes, ...], rIds: [...]}, ...]
    """
    doc = Document(str(filepath))

    text_paras = []
    img_paras = []

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        blips = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        embeds = [b.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                  for b in blips if b.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')]
        if text:
            text_paras.append((i, text))
        if embeds:
            img_paras.append((i, embeds))

    products = []
    for ti, (t_idx, text) in enumerate(text_paras):
        next_t_idx = text_paras[ti + 1][0] if ti + 1 < len(text_paras) else len(doc.paragraphs)
        rids = []
        for i_idx, embeds in img_paras:
            if t_idx < i_idx < next_t_idx:
                rids.extend(embeds)

        blobs = []
        for rid in rids:
            if rid in doc.part.rels:
                rel = doc.part.rels[rid]
                if 'image' in rel.reltype:
                    blobs.append(rel.target_part.blob)

        products.append({
            'name': text,
            'rIds': rids,
            'image_blobs': blobs,
        })

    return products


def match_product(docx_name: str, loai: str) -> Product | None:
    """Tìm sản phẩm trong DB khớp với tên từ docx."""
    codes = extract_model_codes(docx_name)
    if not codes:
        return None

    for code in codes:
        qs = Product.objects.filter(loai=loai, is_active=True).filter(
            Q(ten_hang__icontains=code) | Q(model_turbo__icontains=code)
        )
        if qs.exists():
            return qs.first()

    return None


def create_product(docx_name: str, loai: str) -> Product:
    """Tạo sản phẩm mới từ tên docx."""
    brand = detect_brand(docx_name)
    hang_may = get_hang_may(brand)
    ma_vt = get_next_ma_vt()
    codes = extract_model_codes(docx_name)

    # Tạo tên hiển thị đẹp
    if loai == 'ket_nhot':
        display_name = f'Két làm mát nhớt thủy lực {docx_name} - {brand}'
    else:
        display_name = f'Két nước {docx_name} - {brand}'

    product = Product.objects.create(
        ma_vt=ma_vt,
        ten_hang=display_name,
        loai=loai,
        hang_may=hang_may,
        model_turbo=codes[0] if codes else '',
        is_active=True,
    )
    return product


def save_image(image_blob: bytes, filename: str) -> str:
    """Lưu ảnh vào media/products/ và trả về URL relative path."""
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)

    safe_name = re.sub(r'[\\/:*?"<>|]', '-', filename)
    safe_name = safe_name.replace(' ', '-').lower()

    ext = '.png'
    if image_blob[:4] in (b'\xff\xd8\xff\xe0', b'\xff\xd8\xff\xe1'):
        ext = '.jpg'
    elif image_blob[:8] == b'\x89PNG\r\n\x1a\n':
        ext = '.png'
    elif image_blob[:6] in (b'GIF87a', b'GIF89a'):
        ext = '.gif'
    elif image_blob[:2] == b'BM':
        ext = '.bmp'

    filepath = MEDIA_DIR / f'{safe_name}{ext}'
    counter = 1
    while filepath.exists():
        filepath = MEDIA_DIR / f'{safe_name}_{counter}{ext}'
        counter += 1

    filepath.write_bytes(image_blob)
    return f'/media/products/{filepath.name}'


def import_images(file_keyword=None, limit=None, dry_run=False, create_missing=False):
    """Main import function."""
    results = {
        'matched': 0, 'no_match': 0, 'updated': 0,
        'created': 0, 'errors': [],
    }

    files_to_process = [file_keyword] if file_keyword else list(FILE_LOAI_MAP.keys())

    for fkey in files_to_process:
        filepath = DOCS_DIR / f'{fkey}.docx'
        if not filepath.exists():
            print(f'[ERROR] File khong ton tai: {filepath}')
            continue

        loai = FILE_LOAI_MAP[fkey]
        sep = '=' * 60
        print(f'\n{sep}')
        print(f'[FILE] {fkey}.docx -> loai={loai}')
        print(sep)

        products = parse_docx(filepath)
        print(f'  Tong san pham trong file: {len(products)}')
        items = products[:limit] if limit else products

        for i, prod in enumerate(items):
            name = prod['name']
            codes = extract_model_codes(name)
            n_imgs = len(prod['image_blobs'])

            print(f'\n  [{i+1}/{len(items)}] "{name}"')
            print(f'       Ma model: {codes}')
            print(f'       Anh: {n_imgs} file(s)')

            db_product = match_product(name, loai)

            if db_product:
                print(f'       [MATCH] DB: [{db_product.ma_vt}] "{db_product.ten_hang[:60]}"')
                results['matched'] += 1
            elif create_missing and codes:
                brand = detect_brand(name)
                if not dry_run:
                    db_product = create_product(name, loai)
                    print(f'       [CREATED] [{db_product.ma_vt}] brand={brand} hang_may="{db_product.hang_may.ten}"')
                    results['created'] += 1
                else:
                    print(f'       [WOULD CREATE] brand={brand}')
                    results['created'] += 1
            else:
                print(f'       [NO MATCH] Khong tim thay trong DB')
                results['no_match'] += 1
                continue

            # Import images
            if not dry_run and db_product and prod['image_blobs']:
                img_urls = []
                for j, blob in enumerate(prod['image_blobs']):
                    img_name = f'{fkey.lower().replace(" ", "-")}-{codes[0] if codes else "unknown"}-{j}'
                    url = save_image(blob, img_name)
                    img_urls.append(url)
                    print(f'       [IMAGE] {url} ({len(blob)} bytes)')

                db_product.hinh_anh = img_urls[0]
                db_product.save(update_fields=['hinh_anh', 'updated_at'])
                results['updated'] += 1
                print(f'       [SAVED] Da cap nhat hinh_anh!')

    # Summary
    sep = '=' * 60
    print(f'\n{sep}')
    print('TONG KET')
    print(f'   Match thanh cong: {results["matched"]}')
    if create_missing:
        print(f'   Tao moi: {results["created"]}')
    print(f'   Khong match: {results["no_match"]}')
    print(f'   Da cap nhat anh: {results["updated"]}')
    if dry_run:
        print('   [DRY RUN] - Chua ghi thuc te vao DB')
    print(sep)

    return results


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Import hinh anh tu docx vao database')
    parser.add_argument('--dry-run', action='store_true', help='Chi preview, khong ghi DB')
    parser.add_argument('--file', type=str, help='Chi xu ly 1 file: "KET NHOT" hoac "KET NUOC"')
    parser.add_argument('--limit', type=int, default=None, help='So SP demo (default: tat ca)')
    parser.add_argument('--create-missing', action='store_true', help='Tao SP moi neu chua co trong DB')
    args = parser.parse_args()

    import_images(
        file_keyword=args.file,
        limit=args.limit,
        dry_run=args.dry_run,
        create_missing=args.create_missing,
    )
